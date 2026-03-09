#!/usr/bin/env python3
"""
小学托管班菜单生成器 - Streamlit 主程序（老年人简化版）
"""
import sys
from pathlib import Path
from datetime import datetime
from io import BytesIO

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 添加项目路径
sys.path.insert(0, str(Path(__file__).parent))

from scraper import load_recipes, update_recipes_database
from menu_generator import MenuGenerator, MenuConfig, MealConfig


# 页面配置
st.set_page_config(
    page_title="小学托管班菜单生成器",
    page_icon="🍱",
    layout="centered"
)

# 移动端优化样式
st.markdown("""
<style>
    /* 全局样式 - 大字体，适合老年人 */
    .block-container {
        padding: 1.5rem !important;
    }
    /* 标题 */
    .stMarkdown h1 {
        font-size: 1.6rem !important;
        text-align: center !important;
        margin-bottom: 1rem !important;
    }
    .stMarkdown h2 {
        font-size: 1.2rem !important;
        margin-top: 1rem !important;
        margin-bottom: 0.5rem !important;
    }
    /* 按钮 - 大按钮 */
    .stButton button {
        width: 100% !important;
        font-size: 1.2rem !important;
        padding: 1rem !important;
        height: auto !important;
    }
    /* 输入框 - 大字体 */
    .stNumberInput input {
        font-size: 1.1rem !important;
        height: 3rem !important;
    }
    .stNumberInput label {
        font-size: 1rem !important;
    }
    /* 复选框 - 大字体 */
    .stCheckbox {
        font-size: 1.1rem !important;
    }
    /* 表格 - 大字体 */
    .stDataFrame {
        font-size: 0.95rem !important;
    }
    /* 下拉框 */
    .stSelectbox label {
        font-size: 1rem !important;
    }
    /* 分隔线 */
    hr {
        margin: 1.5rem 0 !important;
    }
    /* 移动端优化 */
    @media (max-width: 768px) {
        .stMarkdown h1 {
            font-size: 1.4rem !important;
        }
        .stButton button {
            font-size: 1.1rem !important;
            padding: 0.8rem !important;
        }
        .stDataFrame {
            font-size: 0.85rem !important;
        }
        .stDataFrame [data-testid="stDataFrameResizable"] {
            overflow-x: auto !important;
        }
    }
</style>
""", unsafe_allow_html=True)

# 标题
st.title("🍱 小学托管班菜单生成器")
st.markdown("---")


@st.cache_data
def get_recipes():
    """加载菜谱数据"""
    return load_recipes()


def generate_menu():
    """生成菜单"""
    recipes = get_recipes()
    generator = MenuGenerator(recipes)

    config = MenuConfig()

    # 早餐配置
    config.breakfast.enabled = st.session_state.get("breakfast_enabled", True)

    # 午餐配置
    config.lunch.enabled = st.session_state.get("lunch_enabled", True)
    config.lunch.meat_count = st.session_state.get("lunch_meat", 2)
    config.lunch.veg_count = st.session_state.get("lunch_veg", 2)
    config.lunch.cold_count = st.session_state.get("lunch_cold", 1)
    config.lunch.soup_count = st.session_state.get("lunch_soup", 1)

    # 晚餐配置
    config.dinner.enabled = st.session_state.get("dinner_enabled", True)
    config.dinner.meat_count = st.session_state.get("dinner_meat", 1)
    config.dinner.veg_count = st.session_state.get("dinner_veg", 2)
    config.dinner.soup_count = st.session_state.get("dinner_soup", 1)
    config.dinner.noodle_count = st.session_state.get("dinner_noodle", 1)

    weekly_menu = generator.generate_weekly_menu(config)
    st.session_state["weekly_menu"] = weekly_menu
    st.session_state["ingredients"] = generator.get_all_ingredients(weekly_menu)


def get_dish_category_label(category):
    """获取菜品分类标签"""
    labels = {
        "荤菜": "🥩",
        "素菜": "🥬",
        "凉菜": "🥗",
        "汤类": "🍜",
        "面食": "🍝",
        "主食": "🍚",
        "早餐": "🍳"
    }
    return labels.get(category, "")


def display_menu_expanded():
    """显示展开式菜单（整周一个表格）"""
    if "weekly_menu" not in st.session_state:
        return

    weekly_menu = st.session_state["weekly_menu"]

    st.subheader("📅 本周菜单")

    table_data = []

    for daily in weekly_menu:
        day_label = f"{daily.day}\n{daily.date}"

        # 早餐
        if daily.breakfast:
            for i, dish in enumerate(daily.breakfast):
                ing_str = "、".join(dish.ingredients)
                table_data.append({
                    "日期": day_label if i == 0 else "",
                    "餐次": "早餐" if i == 0 else "",
                    "菜品": f"{get_dish_category_label(dish.category)} {dish.name}",
                    "所需食材": ing_str
                })

        # 午餐
        if daily.lunch:
            for i, dish in enumerate(daily.lunch):
                ing_str = "、".join(dish.ingredients)
                has_prev = (daily.breakfast and len(daily.breakfast) > 0)
                table_data.append({
                    "日期": day_label if (not has_prev and i == 0) else "",
                    "餐次": "午餐" if i == 0 else "",
                    "菜品": f"{get_dish_category_label(dish.category)} {dish.name}",
                    "所需食材": ing_str
                })

        # 晚餐
        if daily.dinner:
            for i, dish in enumerate(daily.dinner):
                ing_str = "、".join(dish.ingredients)
                has_prev = (daily.breakfast and len(daily.breakfast) > 0) or (daily.lunch and len(daily.lunch) > 0)
                table_data.append({
                    "日期": day_label if (not has_prev and i == 0) else "",
                    "餐次": "晚餐" if i == 0 else "",
                    "菜品": f"{get_dish_category_label(dish.category)} {dish.name}",
                    "所需食材": ing_str
                })

    if table_data:
        df = pd.DataFrame(table_data)
        st.dataframe(
            df,
            hide_index=True,
            use_container_width=True,
            column_config={
                "日期": st.column_config.TextColumn(width="medium"),
                "餐次": st.column_config.TextColumn(width="small"),
                "菜品": st.column_config.TextColumn(width="medium"),
                "所需食材": st.column_config.TextColumn(width="large"),
            }
        )


def export_to_word():
    """导出到 Word 文档（整周一个表格）"""
    if "weekly_menu" not in st.session_state:
        return

    weekly_menu = st.session_state["weekly_menu"]

    doc = Document()

    # 设置默认字体
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
    doc.styles['Normal'].font.size = Pt(11)

    # 标题
    title = doc.add_heading('小学托管班每周菜单', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 日期范围
    if weekly_menu:
        start_date = weekly_menu[0].date
        end_date = weekly_menu[-1].date
        date_para = doc.add_paragraph(f"日期范围：{start_date} 至 {end_date}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 收集整周所有菜品
    all_rows = []
    day_merge_info = []
    meal_merge_info = []

    for daily in weekly_menu:
        day_label = f"{daily.day} ({daily.date})"
        day_start_row = len(all_rows) + 1

        # 早餐
        if daily.breakfast:
            meal_start_row = len(all_rows) + 1
            for i, dish in enumerate(daily.breakfast):
                all_rows.append((
                    day_label if i == 0 else "",
                    "早餐" if i == 0 else "",
                    dish.name,
                    "、".join(dish.ingredients)
                ))
            if len(daily.breakfast) > 1:
                meal_merge_info.append((meal_start_row, len(all_rows), 1))

        # 午餐
        if daily.lunch:
            meal_start_row = len(all_rows) + 1
            has_prev = (daily.breakfast and len(daily.breakfast) > 0)
            for i, dish in enumerate(daily.lunch):
                all_rows.append((
                    day_label if (not has_prev and i == 0) else "",
                    "午餐" if i == 0 else "",
                    dish.name,
                    "、".join(dish.ingredients)
                ))
            if len(daily.lunch) > 1:
                meal_merge_info.append((meal_start_row, len(all_rows), 1))

        # 晚餐
        if daily.dinner:
            meal_start_row = len(all_rows) + 1
            has_prev = (daily.breakfast and len(daily.breakfast) > 0) or (daily.lunch and len(daily.lunch) > 0)
            for i, dish in enumerate(daily.dinner):
                all_rows.append((
                    day_label if (not has_prev and i == 0) else "",
                    "晚餐" if i == 0 else "",
                    dish.name,
                    "、".join(dish.ingredients)
                ))
            if len(daily.dinner) > 1:
                meal_merge_info.append((meal_start_row, len(all_rows), 1))

        day_end_row = len(all_rows)
        if day_end_row > day_start_row:
            day_merge_info.append((day_start_row, day_end_row, 0))

    if all_rows:
        table = doc.add_table(rows=len(all_rows) + 1, cols=4)
        table.style = 'Table Grid'
        table.autofit = True

        for row in table.rows:
            row.cells[0].width = Cm(3)
            row.cells[1].width = Cm(2)
            row.cells[2].width = Cm(5)
            row.cells[3].width = Cm(9)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '日期'
        hdr_cells[1].text = '餐次'
        hdr_cells[2].text = '菜品'
        hdr_cells[3].text = '所需食材'

        for cell in hdr_cells:
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, (day, meal, dish, ing) in enumerate(all_rows):
            row_cells = table.rows[i + 1].cells
            row_cells[0].text = day
            row_cells[1].text = meal
            row_cells[2].text = dish
            row_cells[3].text = ing

            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10)

        for start, end, col_idx in day_merge_info:
            if end > start:
                table.rows[start].cells[col_idx].merge(table.rows[end - 1].cells[col_idx])
                merged = table.rows[start].cells[col_idx]
                merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for start, end, col_idx in meal_merge_info:
            if end > start:
                table.rows[start].cells[col_idx].merge(table.rows[end - 1].cells[col_idx])
                merged = table.rows[start].cells[col_idx]
                merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    st.download_button(
        label="📥 下载 Word 文档",
        data=output,
        file_name=f"每周菜单_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


def export_to_excel():
    """导出到 Excel（整周一个表格）"""
    if "weekly_menu" not in st.session_state:
        return

    weekly_menu = st.session_state["weekly_menu"]

    output = BytesIO()

    table_data = []

    for daily in weekly_menu:
        day_label = f"{daily.day} ({daily.date})"

        if daily.breakfast:
            for i, dish in enumerate(daily.breakfast):
                table_data.append({
                    "日期": day_label if i == 0 else "",
                    "餐次": "早餐" if i == 0 else "",
                    "菜品": dish.name,
                    "所需食材": "、".join(dish.ingredients)
                })

        if daily.lunch:
            for i, dish in enumerate(daily.lunch):
                has_prev = (daily.breakfast and len(daily.breakfast) > 0)
                table_data.append({
                    "日期": day_label if (not has_prev and i == 0) else "",
                    "餐次": "午餐" if i == 0 else "",
                    "菜品": dish.name,
                    "所需食材": "、".join(dish.ingredients)
                })

        if daily.dinner:
            for i, dish in enumerate(daily.dinner):
                has_prev = (daily.breakfast and len(daily.breakfast) > 0) or (daily.lunch and len(daily.lunch) > 0)
                table_data.append({
                    "日期": day_label if (not has_prev and i == 0) else "",
                    "餐次": "晚餐" if i == 0 else "",
                    "菜品": dish.name,
                    "所需食材": "、".join(dish.ingredients)
                })

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        if table_data:
            df = pd.DataFrame(table_data)
            df.to_excel(writer, sheet_name="本周菜单", index=False)

    output.seek(0)

    st.download_button(
        label="📥 下载 Excel 文件",
        data=output,
        file_name=f"每周菜单_{datetime.now().strftime('%Y%m%d')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )


# ===== 主界面 - 配置区域 =====
st.subheader("⚙️ 第一步：选择餐次")

col1, col2, col3 = st.columns(3)
with col1:
    st.session_state["breakfast_enabled"] = st.checkbox("早餐", value=True)
with col2:
    st.session_state["lunch_enabled"] = st.checkbox("午餐", value=True)
with col3:
    st.session_state["dinner_enabled"] = st.checkbox("晚餐", value=True)

st.markdown("---")

if st.session_state.get("lunch_enabled", True):
    st.subheader("🍽️ 第二步：午餐配置")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.session_state["lunch_meat"] = st.number_input("荤菜", min_value=0, max_value=5, value=2)
    with col2:
        st.session_state["lunch_veg"] = st.number_input("素菜", min_value=0, max_value=5, value=2)
    with col3:
        st.session_state["lunch_cold"] = st.number_input("凉菜", min_value=0, max_value=5, value=1)
    with col4:
        st.session_state["lunch_soup"] = st.number_input("汤类", min_value=0, max_value=5, value=1)

    st.markdown("---")

if st.session_state.get("dinner_enabled", True):
    st.subheader("🍜 第三步：晚餐配置")
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.session_state["dinner_meat"] = st.number_input("荤菜", min_value=0, max_value=5, value=1, key="dinner_meat_key")
    with col2:
        st.session_state["dinner_veg"] = st.number_input("素菜", min_value=0, max_value=5, value=2, key="dinner_veg_key")
    with col3:
        st.session_state["dinner_soup"] = st.number_input("汤类", min_value=0, max_value=5, value=1, key="dinner_soup_key")
    with col4:
        st.session_state["dinner_noodle"] = st.number_input("面食", min_value=0, max_value=5, value=1, key="dinner_noodle_key")

    st.markdown("---")

# 生成按钮 - 居中放上面
st.subheader("✨ 第四步：生成菜单")
if st.button("🎉 生成本周菜单", type="primary"):
    with st.spinner("正在生成菜单，请稍候..."):
        generate_menu()
    st.success("✅ 菜单生成成功！")
    st.balloons()

st.markdown("---")

# 显示菜单
if "weekly_menu" in st.session_state:
    display_menu_expanded()

    st.markdown("---")

    # 导出按钮
    st.subheader("📤 第五步：下载菜单")
    col1, col2 = st.columns(2)
    with col1:
        export_to_word()
    with col2:
        export_to_excel()

else:
    st.info("👆 请先在上面配置好，然后点击「生成本周菜单」按钮！")


st.markdown("---")
st.caption("小学托管班菜单生成器")
